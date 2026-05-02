from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_PATH = "/Users/lns2/Desktop/lnsapp-ios/ios_backend_subscription_execution_plan.docx"
MARKER_NEW = "本次新增-重启容灾"
MARKER_MOD = "本次修改-重启容灾"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(17, 59, 109)
    p.space_after = Pt(8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(90, 90, 90)
    p2.space_after = Pt(16)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        r2.bold = False
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_marker_note(doc, kind, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"【{kind}】")
    r1.bold = True
    if "新增" in kind:
        r1.font.color.rgb = RGBColor(192, 57, 43)
    else:
        r1.font.color.rgb = RGBColor(211, 84, 0)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(70, 70, 70)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True)
        shade_cell(hdr.cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Menlo"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
        r.font.size = Pt(8.5)


def table_rows_for_fields(fields):
    return [[
        f["name"],
        f["type"],
        f["required"],
        f["source"],
        f["when"],
        (f"【{f['marker']}】" if f.get("marker") else "") + f["comment"],
    ] for f in fields]


def add_field_table(doc, title, fields):
    add_heading(doc, title, level=3)
    add_table(
        doc,
        ["字段", "类型", "必填", "来源", "写入/返回时机", "注释"],
        table_rows_for_fields(fields),
    )


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)

    add_title(
        doc,
        "Elavatine iOS 订阅后台研发执行方案",
        "适用范围：iOS App Store 自动续订订阅 | 面向对象：后端研发 / 联调研发 / 运维排障 | 含服务器重启与短暂不可用容灾补充",
    )

    add_heading(doc, "0. 本次修订标记说明", level=1)
    add_marker_note(doc, MARKER_NEW, "表示这次因为“后台服务器偶尔重启、首购请求可能收不到、Apple续费通知可能漏接”而新增的设计内容。")
    add_marker_note(doc, MARKER_MOD, "表示原方案已有，但这次因为重启容灾要求而进一步增强或调整的内容。")
    add_para(doc, "阅读建议：后台同学可优先搜索文档中的“【本次新增-重启容灾】”和“【本次修改-重启容灾】”，这样可以快速看到本轮补充的重点。")

    add_heading(doc, "1. 文档目的", level=1)
    add_para(doc, "这份文档用于指导后台研发落地 Elavatine iOS 订阅系统改造。目标不是只把支付接口打通，而是建立一套“交易可追溯、重复提交可幂等、Apple 通知可回补、注册前购买可承接、会员权益可自动补齐”的订阅账本体系。")
    add_para(doc, "为什么要这样做：当前 iOS 端已经存在注册前购买场景（GuidanceProVC），且现有 App 代码仍以 StoreKit 1 与后置绑定为主。如果后台只保留当前三张正式业务表，会出现匿名购买难承接、通知漏单难回补、支付成功但会员未生效难追责等风险。")
    add_marker_note(doc, MARKER_NEW, "本次额外把“后台重启、容器滚动发布、短暂 502/超时、Apple 通知投递时服务不可用”明确纳入设计范围，目标是保证首购和续费都能在恢复后自动补齐。")

    add_heading(doc, "2. 设计原则", level=1)
    add_bullets(doc, [
        "Apple 交易是真相源：transaction_id、original_transaction_id、web_order_line_item_id 代表 Apple 侧真实交易链。",
        "服务端账本是业务真相：所有会员权益最终以服务端计算结果为准，客户端只做触发与恢复。",
        "正式业务表继续保留：t_user_subscription_contract、t_user_subscription_order、t_user_vip_info 继续作为正式用户态表。",
        "注册前购买必须有匿名身份承接：不能等登录后才创建订阅身份。",
        "所有链路必须幂等：客户端重试、通知重复投递、对账任务回补，都不能产生重复开通或错绑。",
        "通知不是唯一恢复路径：客户端恢复购买、服务端主动查 Apple、Notification History 回补都必须存在。",
    ])
    add_marker_note(doc, MARKER_NEW, "新增一条容灾原则：服务端短暂不可用不是例外场景，而是订阅系统的标准故障模型。设计上必须允许客户端重试、Apple 重试、服务端历史回补三条恢复链并存。")

    add_heading(doc, "3. 现有三张正式业务表的定位", level=1)
    add_bullets(doc, [
        "t_user_subscription_contract：正式归属到 uid 的订阅合同主表；iOS 下 contract_id = original_transaction_id。",
        "t_user_subscription_order：正式归属到 uid 的业务订单流水；iOS 下 third_party_trade_no = transaction_id。",
        "t_user_vip_info：App 唯一会员真相表；前端展示和业务权限判断只认这张表。",
    ])
    add_para(doc, "为什么不能只靠这三张表：因为注册前购买时还没有 uid，这三张表都无法自然承接匿名交易；同时 Apple 原始交易、通知原文、对账状态、重试状态也不适合直接塞进正式业务表。")

    add_heading(doc, "4. 后台总体架构", level=1)
    add_code_block(doc, """
Apple
  ├─ StoreKit / App Store
  ├─ App Store Server API
  └─ App Store Server Notifications V2

iOS Client
  ├─ purchase / verified transaction
  ├─ pending transaction queue
  ├─ restore purchase
  └─ identity init / purchase report / bind anonymous

Backend
  ├─ t_ios_identity
  ├─ t_ios_transaction_ledger
  ├─ t_ios_notification_inbox
  ├─ t_ios_outage_window
  ├─ t_ios_purchase_attempt
  ├─ t_user_subscription_contract
  ├─ t_user_subscription_order
  ├─ t_user_vip_info
  ├─ ASSN V2 receiver
  ├─ Apple Server API verifier
  ├─ outage recovery jobs
  └─ reconcile jobs
    """)
    add_para(doc, "职责拆分：前端负责发现交易、恢复交易、上报交易；后台负责验真、记账、绑定、对账、长期会员态输出。")
    add_marker_note(doc, MARKER_MOD, "架构层新增 outage recovery jobs 和 t_ios_outage_window，用来显式处理‘服务重启窗口内首购请求或Apple通知漏收’的问题，而不是事后临时手工补。")

    add_heading(doc, "5. 数据表设计", level=1)
    add_para(doc, "本节先写每张表为什么需要，再给出字段级说明。字段说明中的“来源”用于告诉研发该字段从哪里拿；“写入/返回时机”用于告诉研发什么时候填。")

    add_heading(doc, "5.1 新表：t_ios_identity", level=2)
    add_para(doc, "用途：承接注册前购买身份。用户第一次进入注册前付费墙前，后台为客户端签发 anonymous_user_id 和 app_account_token。后续注册成功后再把匿名身份绑定到真实 uid。")
    add_field_table(doc, "t_ios_identity 字段说明", [
        {"name":"id","type":"BIGINT UNSIGNED","required":"是","source":"DB","when":"insert","comment":"主键，自增 ID。"},
        {"name":"anonymous_user_id","type":"VARCHAR(64)","required":"是","source":"后台生成 UUID","when":"identity/init 时写入","comment":"匿名购买主体 ID。注册前购买、恢复购买、人工补单都可用它查到一组匿名交易。"},
        {"name":"app_account_token","type":"CHAR(36)","required":"是","source":"后台生成 UUID","when":"identity/init 时写入","comment":"购买时传给 Apple 的稳定身份 token。Apple 交易、续费、通知、对账都依赖它关联到账户。"},
        {"name":"device_install_id","type":"VARCHAR(64)","required":"是","source":"前端本地生成 UUID","when":"identity/init 时写入","comment":"设备安装实例 ID。用于排查，不作为最终归属依据。"},
        {"name":"uid","type":"VARCHAR(32)","required":"否","source":"登录后用户","when":"bind-anonymous 成功时写入","comment":"真实业务用户 ID。匿名阶段为空。"},
        {"name":"bind_status","type":"TINYINT","required":"是","source":"服务端状态机","when":"init/绑定/冲突时更新","comment":"1 anonymous，2 bound，3 conflict，4 disabled。用于标记匿名身份是否已绑定真实账号。"},
        {"name":"created_at","type":"DATETIME","required":"是","source":"DB","when":"insert","comment":"创建时间。"},
        {"name":"updated_at","type":"DATETIME","required":"是","source":"DB","when":"update","comment":"更新时间。"},
    ])

    add_heading(doc, "5.2 新表：t_ios_transaction_ledger", level=2)
    add_para(doc, "用途：Apple 原始交易账本。它不是给前端看的业务表，而是后台长期真相表。所有 Apple 客户端上报、Apple 通知、对账回补，最终都先落到这张表。")
    add_marker_note(doc, MARKER_MOD, "账本表本次补了“正式落账状态、最近一次服务端核验时间、最近一次同步来源”三个字段，用来支撑后台重启后的自动补齐与查漏。")
    add_field_table(doc, "t_ios_transaction_ledger 字段说明", [
        {"name":"id","type":"BIGINT UNSIGNED","required":"是","source":"DB","when":"insert","comment":"主键。"},
        {"name":"transaction_id","type":"VARCHAR(128)","required":"是","source":"Apple transaction.id / transactionId","when":"verified 上报、通知、对账时写入","comment":"Apple 单笔交易唯一 ID。整张账本最重要的幂等键之一。必须唯一。"},
        {"name":"original_transaction_id","type":"VARCHAR(128)","required":"是","source":"Apple originalTransactionId","when":"verified 上报、通知、对账时写入","comment":"自动续订链主键。用于把首购、续费、恢复、退款串成一条链。"},
        {"name":"web_order_line_item_id","type":"VARCHAR(128)","required":"否","source":"Apple transaction payload","when":"可取到时写入","comment":"Apple 续费行项目 ID。排查续费链和账期切换时非常有用。"},
        {"name":"product_id","type":"VARCHAR(64)","required":"是","source":"Apple productID","when":"交易入账时写入","comment":"购买的商品 ID。"},
        {"name":"app_account_token","type":"CHAR(36)","required":"否","source":"Apple payload / t_ios_identity","when":"交易入账时写入","comment":"交易对应的购买身份 token。用于把匿名身份与 Apple 交易链对齐。"},
        {"name":"anonymous_user_id","type":"VARCHAR(64)","required":"否","source":"t_ios_identity","when":"匿名购买时写入","comment":"若该交易来自注册前购买，记录匿名主体。"},
        {"name":"uid","type":"VARCHAR(32)","required":"否","source":"绑定后的真实用户","when":"绑定成功或已知用户时写入","comment":"正式用户 ID。用于后续对账、退款、客服查询。"},
        {"name":"environment","type":"VARCHAR(16)","required":"是","source":"Apple","when":"交易入账时写入","comment":"Sandbox 或 Production。沙盒和正式必须严格区分。"},
        {"name":"signed_transaction_info","type":"MEDIUMTEXT","required":"否","source":"客户端上报 / Apple API","when":"客户端 verified 上报或服务端查交易时写入","comment":"原始 signedTransactionInfo JWS，用于验签和审计。"},
        {"name":"signed_renewal_info","type":"MEDIUMTEXT","required":"否","source":"Apple 通知 / Apple API","when":"收到 renewal 信息时写入","comment":"原始 signedRenewalInfo JWS，用于判断自动续费状态、宽限期、账单重试等。"},
        {"name":"ledger_status","type":"VARCHAR(32)","required":"是","source":"服务端计算","when":"每次状态变化时更新","comment":"交易账本状态，如 verified、refunded、revoked、expired、grace、billing_retry。"},
        {"name":"formal_status","type":"VARCHAR(32)","required":"是","source":"服务端计算","when":"每次补齐正式三表后更新","comment":"正式落账状态：anonymous_pending_bind/formalized/reconcile_pending/reconcile_failed。用于判断交易是否已经成功映射到正式业务表。","marker":MARKER_NEW},
        {"name":"last_verified_at","type":"DATETIME","required":"否","source":"服务端当前时间","when":"每次成功核验 Apple 后更新","comment":"最近一次完成 Apple 核验的时间。后台重启恢复后做对账时优先看这个时间。","marker":MARKER_NEW},
        {"name":"last_sync_source","type":"VARCHAR(24)","required":"否","source":"服务端固定","when":"每次状态变更时更新","comment":"最近一次更新来源：client_report/notification/reconcile/restore/manual。用于排查到底是哪条链修复了这笔交易。","marker":MARKER_NEW},
        {"name":"source","type":"VARCHAR(16)","required":"是","source":"服务端固定","when":"insert","comment":"client_report / notification / reconcile / manual，表示该账本记录从哪条链路进入。"},
        {"name":"created_at","type":"DATETIME","required":"是","source":"DB","when":"insert","comment":"创建时间。"},
        {"name":"updated_at","type":"DATETIME","required":"是","source":"DB","when":"update","comment":"更新时间。"},
    ])

    add_heading(doc, "5.3 新表：t_ios_notification_inbox", level=2)
    add_para(doc, "用途：Apple Server Notifications V2 收件箱。不要收到通知就直接改正式业务表，必须先落库，再异步处理。")
    add_marker_note(doc, MARKER_MOD, "通知收件箱本次补了“是否历史回放、首次发送结果、关联故障窗口”三个字段，用于恢复后台重启期间漏收的通知。")
    add_field_table(doc, "t_ios_notification_inbox 字段说明", [
        {"name":"id","type":"BIGINT UNSIGNED","required":"是","source":"DB","when":"insert","comment":"主键。"},
        {"name":"notification_uuid","type":"VARCHAR(64)","required":"是","source":"Apple notificationUUID","when":"回调时写入","comment":"Apple 通知唯一 ID。用于去重，必须唯一。"},
        {"name":"notification_type","type":"VARCHAR(64)","required":"是","source":"Apple payload","when":"回调时写入","comment":"通知类型，如 SUBSCRIBED、DID_RENEW、EXPIRED、REFUND 等。"},
        {"name":"subtype","type":"VARCHAR(64)","required":"否","source":"Apple payload","when":"回调时写入","comment":"通知子类型。不同 subtype 会影响具体处理逻辑。"},
        {"name":"original_transaction_id","type":"VARCHAR(128)","required":"否","source":"Apple payload decodedTransactionInfo","when":"可取到时写入","comment":"通知关联的订阅链。"},
        {"name":"signed_payload","type":"MEDIUMTEXT","required":"是","source":"Apple signedPayload","when":"回调时写入","comment":"原始通知 JWS。必须原样留档，便于审计与回放。"},
        {"name":"process_status","type":"VARCHAR(24)","required":"是","source":"服务端状态机","when":"接收后为 pending，处理后更新","comment":"pending / success / failed / ignored。"},
        {"name":"process_error","type":"VARCHAR(255)","required":"否","source":"服务端异常信息","when":"处理失败时更新","comment":"最近一次处理失败原因。"},
        {"name":"is_replayed","type":"TINYINT","required":"是","source":"服务端","when":"Notification History 回补入库时写入","comment":"0 正常实时通知，1 表示通过 Get Notification History 回放得到。用于区分实时接收与事后补回。","marker":MARKER_NEW},
        {"name":"first_send_attempt_result","type":"VARCHAR(64)","required":"否","source":"Apple Notification History","when":"通过 Notification History 回补时写入","comment":"Apple 首次投递结果，如 SUCCESS/NO_RESPONSE/TIMED_OUT。用于确认是否在服务重启窗口漏收。","marker":MARKER_NEW},
        {"name":"outage_window_id","type":"BIGINT UNSIGNED","required":"否","source":"t_ios_outage_window","when":"识别出该通知属于某次故障窗口时写入","comment":"关联的故障窗口记录，方便一次性回放和审计。","marker":MARKER_NEW},
        {"name":"arrived_at","type":"DATETIME","required":"是","source":"服务端当前时间","when":"回调时写入","comment":"通知到达时间。"},
        {"name":"processed_at","type":"DATETIME","required":"否","source":"服务端当前时间","when":"处理成功或失败时写入","comment":"处理完成时间。"},
    ])

    add_heading(doc, "5.4 新表：t_ios_outage_window", level=2)
    add_para(doc, "用途：记录后台服务器重启、滚动发布、健康检查失败、通知入口不可用等故障窗口。恢复任务会按这个窗口去回补 Notification History 和未正式入账交易。")
    add_marker_note(doc, MARKER_NEW, "这是本次专门为了“后台服务器偶尔会重启”补进去的新表。没有这张表，恢复任务很难明确知道需要补哪一段时间的通知和交易。")
    add_field_table(doc, "t_ios_outage_window 字段说明", [
        {"name":"id","type":"BIGINT UNSIGNED","required":"是","source":"DB","when":"insert","comment":"主键。"},
        {"name":"outage_key","type":"VARCHAR(64)","required":"是","source":"服务端生成","when":"检测到故障窗口时写入","comment":"故障窗口唯一键，可按日期+实例+类型组合生成。"},
        {"name":"started_at","type":"DATETIME","required":"是","source":"运维/健康检查/服务端","when":"故障开始时写入","comment":"故障开始时间。用于决定回补 Notification History 和交易对账的起点。"},
        {"name":"ended_at","type":"DATETIME","required":"否","source":"运维/健康检查/服务端","when":"恢复后写入","comment":"故障结束时间。"},
        {"name":"outage_type","type":"VARCHAR(24)","required":"是","source":"服务端固定","when":"insert","comment":"restart/deploy/unhealthy/notification_down/api_down 等。"},
        {"name":"affected_scope","type":"VARCHAR(24)","required":"是","source":"服务端固定","when":"insert","comment":"api/notifications/all，表示影响范围。"},
        {"name":"recover_status","type":"VARCHAR(24)","required":"是","source":"服务端状态机","when":"恢复任务执行时更新","comment":"pending/recovering/done/failed。"},
        {"name":"last_recover_at","type":"DATETIME","required":"否","source":"服务端当前时间","when":"每次恢复任务运行时更新","comment":"最近一次跑恢复任务的时间。"},
        {"name":"remark","type":"VARCHAR(255)","required":"否","source":"运维/服务端","when":"需要时写入","comment":"补充说明，如实例编号、部署单号、异常原因。"},
        {"name":"created_at","type":"DATETIME","required":"是","source":"DB","when":"insert","comment":"创建时间。"},
        {"name":"updated_at","type":"DATETIME","required":"是","source":"DB","when":"update","comment":"更新时间。"},
    ])

    add_heading(doc, "5.5 新表：t_ios_purchase_attempt", level=2)
    add_para(doc, "用途：记录客户端一次购买尝试以及它的重试和同步状态。它不是正式订单，而是客户端购买链的过程账。")
    add_marker_note(doc, MARKER_MOD, "购买尝试表本次补了“服务端是否已确认持久化、最近一次服务端确认时间、下次建议重试时间”，用于处理后台重启后客户端请求无响应的情况。")
    add_field_table(doc, "t_ios_purchase_attempt 字段说明", [
        {"name":"id","type":"BIGINT UNSIGNED","required":"是","source":"DB","when":"insert","comment":"主键。"},
        {"name":"attempt_id","type":"VARCHAR(64)","required":"是","source":"前端 UUID","when":"点击购买前生成","comment":"一次购买尝试 ID。客户端重复重试、后台排查时可用。"},
        {"name":"anonymous_user_id","type":"VARCHAR(64)","required":"否","source":"t_ios_identity","when":"注册前购买时写入","comment":"匿名购买主体。"},
        {"name":"app_account_token","type":"CHAR(36)","required":"否","source":"t_ios_identity","when":"购买发起时写入","comment":"购买身份 token。"},
        {"name":"uid","type":"VARCHAR(32)","required":"否","source":"登录用户","when":"购买时已登录或绑定后写入","comment":"真实业务用户 ID。"},
        {"name":"transaction_id","type":"VARCHAR(128)","required":"否","source":"Apple","when":"verified 后写入","comment":"若购买已被 Apple 验证，则记录交易 ID。"},
        {"name":"original_transaction_id","type":"VARCHAR(128)","required":"否","source":"Apple","when":"verified 后写入","comment":"订阅链 ID。"},
        {"name":"product_id","type":"VARCHAR(64)","required":"是","source":"前端待购商品","when":"点击购买前写入","comment":"商品 ID。"},
        {"name":"client_status","type":"VARCHAR(24)","required":"是","source":"前端","when":"状态变化时更新","comment":"purchasing / verified_local / pending_sync / synced / failed_retryable 等。"},
        {"name":"server_status","type":"VARCHAR(24)","required":"否","source":"服务端","when":"收到上报后更新","comment":"accepted / verified / rejected / retryable_failed 等。"},
        {"name":"server_ack_status","type":"VARCHAR(24)","required":"否","source":"服务端","when":"请求写库完成后更新","comment":"none/persisted/response_lost。用于区分‘后台没处理到’和‘后台已写库但响应在重启窗口里丢了’。","marker":MARKER_NEW},
        {"name":"server_ack_at","type":"DATETIME","required":"否","source":"服务端当前时间","when":"后台确认写库成功时更新","comment":"最近一次确认服务端已持久化的时间。","marker":MARKER_NEW},
        {"name":"next_retry_after","type":"DATETIME","required":"否","source":"服务端","when":"判定需要客户端稍后重试时更新","comment":"建议客户端下次重试时间。可用于后台刚恢复时控制重试洪峰。","marker":MARKER_NEW},
        {"name":"retry_count","type":"INT","required":"是","source":"前后端","when":"每次重试时更新","comment":"累计重试次数。"},
        {"name":"last_retry_at","type":"DATETIME","required":"否","source":"前后端","when":"重试时更新","comment":"最近一次重试时间。"},
        {"name":"created_at","type":"DATETIME","required":"是","source":"DB","when":"insert","comment":"创建时间。"},
        {"name":"updated_at","type":"DATETIME","required":"是","source":"DB","when":"update","comment":"更新时间。"},
    ])

    add_heading(doc, "5.6 现有正式表扩字段", level=2)
    add_para(doc, "这三张表继续保留，不做职责替换。新增字段的目标是让它们和 Apple 账本打通，而不是让它们替代 Apple 账本。")

    add_field_table(doc, "t_user_subscription_contract 新增字段说明", [
        {"name":"app_account_token","type":"CHAR(36)","required":"否","source":"t_ios_identity / Apple payload","when":"首次正式绑定或续费到账时写入","comment":"这条正式订阅链对应的身份 token。"},
        {"name":"anonymous_user_id","type":"VARCHAR(64)","required":"否","source":"t_ios_identity","when":"注册前购买成功绑定后写入","comment":"表明这条订阅链来源于匿名购买。"},
        {"name":"latest_transaction_id","type":"VARCHAR(128)","required":"否","source":"Apple","when":"每次新交易入账时更新","comment":"最近一笔交易 ID。"},
        {"name":"environment","type":"VARCHAR(16)","required":"否","source":"Apple","when":"首次建链时写入","comment":"Sandbox 或 Production。"},
        {"name":"last_notification_uuid","type":"VARCHAR(64)","required":"否","source":"t_ios_notification_inbox","when":"通知成功驱动状态变化后更新","comment":"最近一次导致合同状态变化的 Apple 通知。"},
        {"name":"last_verify_time","type":"DATETIME","required":"否","source":"服务端当前时间","when":"每次核验 Apple 成功后更新","comment":"最近一次校验时间。"},
        {"name":"last_error_msg","type":"VARCHAR(255)","required":"否","source":"服务端异常","when":"异常时更新","comment":"最近一次错误原因。"},
        {"name":"recover_window_id","type":"BIGINT UNSIGNED","required":"否","source":"t_ios_outage_window","when":"由故障恢复任务补齐该合同时写入","comment":"标记该正式合同最近一次是由哪次故障恢复任务修复完成。","marker":MARKER_NEW},
    ])

    add_field_table(doc, "t_user_subscription_order 新增字段说明", [
        {"name":"original_transaction_id","type":"VARCHAR(128)","required":"否","source":"Apple","when":"订单入账时写入","comment":"冗余存储订阅链 ID，便于和合同表/ledger 关联查询。"},
        {"name":"web_order_line_item_id","type":"VARCHAR(128)","required":"否","source":"Apple","when":"可取到时写入","comment":"Apple 续费行项目 ID。"},
        {"name":"app_account_token","type":"CHAR(36)","required":"否","source":"t_ios_identity / Apple payload","when":"订单入账时写入","comment":"购买身份 token。"},
        {"name":"anonymous_user_id","type":"VARCHAR(64)","required":"否","source":"t_ios_identity","when":"注册前购买绑定时写入","comment":"匿名主体。"},
        {"name":"source","type":"VARCHAR(16)","required":"否","source":"服务端固定","when":"订单入账时写入","comment":"client_report / notification / reconcile / manual。"},
        {"name":"verify_status","type":"VARCHAR(24)","required":"否","source":"服务端","when":"验 Apple 完成后更新","comment":"verified / unverified / retryable_failed。"},
        {"name":"refund_reason","type":"VARCHAR(128)","required":"否","source":"Apple / 客服","when":"退款时写入","comment":"退款原因说明。"},
        {"name":"recover_window_id","type":"BIGINT UNSIGNED","required":"否","source":"t_ios_outage_window","when":"该订单由故障恢复任务补齐时写入","comment":"标记该业务订单最近一次由哪次故障恢复任务补入。","marker":MARKER_NEW},
    ])

    add_field_table(doc, "t_user_vip_info 新增字段说明", [
        {"name":"source","type":"VARCHAR(16)","required":"否","source":"服务端固定","when":"更新会员态时写入","comment":"ios / wx / alipay / manual。"},
        {"name":"source_contract_id","type":"VARCHAR(128)","required":"否","source":"t_user_subscription_contract.contract_id","when":"更新会员态时写入","comment":"当前生效权益对应的订阅链。"},
        {"name":"sync_status","type":"VARCHAR(24)","required":"否","source":"服务端","when":"同步中或失败时更新","comment":"synced / pending / failed。前端可以据此显示“订单同步中”。"},
        {"name":"grace_status","type":"VARCHAR(24)","required":"否","source":"Apple renewal status","when":"宽限期或账单重试时更新","comment":"none / grace / billing_retry。"},
        {"name":"last_verify_time","type":"DATETIME","required":"否","source":"服务端当前时间","when":"每次会员态核验后更新","comment":"最近一次确认会员态时间。"},
        {"name":"remark","type":"VARCHAR(255)","required":"否","source":"服务端 / 客服","when":"异常补偿时更新","comment":"人工补单、异常说明、备注。"},
        {"name":"recover_window_id","type":"BIGINT UNSIGNED","required":"否","source":"t_ios_outage_window","when":"该会员态由故障恢复任务修正时写入","comment":"用于追踪某次会员修正是否来自服务器重启后的恢复任务。","marker":MARKER_NEW},
    ])

    add_heading(doc, "6. SQL 方案", level=1)
    add_para(doc, "以下 SQL 以 MySQL 8 为假设，可由后台同学再按实际库规范微调。SQL 中已带字段注释。")

    add_heading(doc, "6.1 新表 SQL", level=2)
    add_marker_note(doc, MARKER_MOD, "本次 SQL 部分新增了 t_ios_outage_window，并给 ledger / inbox / purchase_attempt 增加了容灾字段。后台同学如果已经基于上一版开始建表，请重点对比这里。")
    add_code_block(doc, """
CREATE TABLE t_ios_identity (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '主键',
    anonymous_user_id VARCHAR(64) NOT NULL COMMENT '匿名购买主体ID，注册前购买链路使用',
    app_account_token CHAR(36) NOT NULL COMMENT '购买时传给Apple的appAccountToken，稳定UUID',
    device_install_id VARCHAR(64) NOT NULL COMMENT '设备安装实例ID，仅用于排查，不作为最终归属依据',
    uid VARCHAR(32) DEFAULT NULL COMMENT '真实业务用户ID，匿名阶段为空',
    bind_status TINYINT UNSIGNED NOT NULL DEFAULT 1 COMMENT '绑定状态：1-anonymous，2-bound，3-conflict，4-disabled',
    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP COMMENT '更新时间',
    PRIMARY KEY (id),
    UNIQUE KEY uk_anonymous_user_id (anonymous_user_id),
    UNIQUE KEY uk_app_account_token (app_account_token),
    KEY idx_uid (uid),
    KEY idx_device_install_id (device_install_id)
) COMMENT='iOS注册前购买身份承接表';

CREATE TABLE t_ios_transaction_ledger (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '主键',
    transaction_id VARCHAR(128) NOT NULL COMMENT 'Apple单笔交易ID，幂等主键',
    original_transaction_id VARCHAR(128) NOT NULL COMMENT 'Apple自动续订链主键',
    web_order_line_item_id VARCHAR(128) DEFAULT NULL COMMENT 'Apple续费行项目ID',
    product_id VARCHAR(64) NOT NULL COMMENT '商品ID',
    app_account_token CHAR(36) DEFAULT NULL COMMENT '购买身份token',
    anonymous_user_id VARCHAR(64) DEFAULT NULL COMMENT '匿名购买主体ID',
    uid VARCHAR(32) DEFAULT NULL COMMENT '正式业务用户ID',
    environment VARCHAR(16) NOT NULL COMMENT 'Sandbox或Production',
    signed_transaction_info MEDIUMTEXT COMMENT '原始signedTransactionInfo JWS',
    signed_renewal_info MEDIUMTEXT COMMENT '原始signedRenewalInfo JWS',
    ledger_status VARCHAR(32) NOT NULL COMMENT '账本状态：verified/refunded/revoked/expired/grace/billing_retry等',
    formal_status VARCHAR(32) NOT NULL COMMENT '正式落账状态：anonymous_pending_bind/formalized/reconcile_pending/reconcile_failed',
    last_verified_at DATETIME DEFAULT NULL COMMENT '最近一次完成Apple核验的时间',
    last_sync_source VARCHAR(24) DEFAULT NULL COMMENT '最近一次更新来源：client_report/notification/reconcile/restore/manual',
    source VARCHAR(16) NOT NULL COMMENT '来源：client_report/notification/reconcile/manual',
    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP COMMENT '更新时间',
    PRIMARY KEY (id),
    UNIQUE KEY uk_transaction_id (transaction_id),
    KEY idx_original_transaction_id (original_transaction_id),
    KEY idx_web_order_line_item_id (web_order_line_item_id),
    KEY idx_uid (uid),
    KEY idx_app_account_token (app_account_token),
    KEY idx_ledger_status (ledger_status)
) COMMENT='iOS Apple原始交易账本';

CREATE TABLE t_ios_notification_inbox (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '主键',
    notification_uuid VARCHAR(64) NOT NULL COMMENT 'Apple通知唯一ID，去重主键',
    notification_type VARCHAR(64) NOT NULL COMMENT 'Apple通知类型',
    subtype VARCHAR(64) DEFAULT NULL COMMENT 'Apple通知子类型',
    original_transaction_id VARCHAR(128) DEFAULT NULL COMMENT '通知关联的订阅链ID',
    signed_payload MEDIUMTEXT NOT NULL COMMENT '原始通知signedPayload JWS',
    process_status VARCHAR(24) NOT NULL DEFAULT 'pending' COMMENT '处理状态：pending/success/failed/ignored',
    process_error VARCHAR(255) DEFAULT NULL COMMENT '处理失败原因',
    is_replayed TINYINT UNSIGNED NOT NULL DEFAULT 0 COMMENT '是否通过Notification History回放得到：0否，1是',
    first_send_attempt_result VARCHAR(64) DEFAULT NULL COMMENT 'Apple首次投递结果：SUCCESS/NO_RESPONSE/TIMED_OUT等',
    outage_window_id BIGINT UNSIGNED DEFAULT NULL COMMENT '关联的故障恢复窗口ID',
    arrived_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '通知到达时间',
    processed_at DATETIME DEFAULT NULL COMMENT '通知处理完成时间',
    PRIMARY KEY (id),
    UNIQUE KEY uk_notification_uuid (notification_uuid),
    KEY idx_original_transaction_id (original_transaction_id),
    KEY idx_process_status (process_status),
    KEY idx_notification_type (notification_type)
) COMMENT='iOS Apple通知收件箱';

CREATE TABLE t_ios_outage_window (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '主键',
    outage_key VARCHAR(64) NOT NULL COMMENT '故障窗口唯一键',
    started_at DATETIME NOT NULL COMMENT '故障开始时间',
    ended_at DATETIME DEFAULT NULL COMMENT '故障结束时间',
    outage_type VARCHAR(24) NOT NULL COMMENT '故障类型：restart/deploy/unhealthy/notification_down/api_down',
    affected_scope VARCHAR(24) NOT NULL COMMENT '影响范围：api/notifications/all',
    recover_status VARCHAR(24) NOT NULL DEFAULT 'pending' COMMENT '恢复状态：pending/recovering/done/failed',
    last_recover_at DATETIME DEFAULT NULL COMMENT '最近一次执行恢复任务的时间',
    remark VARCHAR(255) DEFAULT NULL COMMENT '补充说明，如实例、发布单号、异常原因',
    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP COMMENT '更新时间',
    PRIMARY KEY (id),
    UNIQUE KEY uk_outage_key (outage_key),
    KEY idx_started_at (started_at),
    KEY idx_recover_status (recover_status),
    KEY idx_affected_scope (affected_scope)
) COMMENT='iOS订阅服务故障窗口记录表';

CREATE TABLE t_ios_purchase_attempt (
    id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT COMMENT '主键',
    attempt_id VARCHAR(64) NOT NULL COMMENT '一次购买尝试ID，前端生成',
    anonymous_user_id VARCHAR(64) DEFAULT NULL COMMENT '匿名购买主体ID',
    app_account_token CHAR(36) DEFAULT NULL COMMENT '购买身份token',
    uid VARCHAR(32) DEFAULT NULL COMMENT '真实业务用户ID',
    transaction_id VARCHAR(128) DEFAULT NULL COMMENT 'Apple单笔交易ID',
    original_transaction_id VARCHAR(128) DEFAULT NULL COMMENT 'Apple自动续订链主键',
    product_id VARCHAR(64) NOT NULL COMMENT '商品ID',
    client_status VARCHAR(24) NOT NULL COMMENT '客户端状态：purchasing/verified_local/pending_sync/synced/failed_retryable等',
    server_status VARCHAR(24) DEFAULT NULL COMMENT '服务端状态：accepted/verified/rejected/retryable_failed等',
    server_ack_status VARCHAR(24) DEFAULT NULL COMMENT '服务端确认状态：none/persisted/response_lost',
    server_ack_at DATETIME DEFAULT NULL COMMENT '最近一次服务端确认写库完成时间',
    next_retry_after DATETIME DEFAULT NULL COMMENT '建议客户端下次重试时间',
    retry_count INT NOT NULL DEFAULT 0 COMMENT '累计重试次数',
    last_retry_at DATETIME DEFAULT NULL COMMENT '最近一次重试时间',
    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '创建时间',
    updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP COMMENT '更新时间',
    PRIMARY KEY (id),
    UNIQUE KEY uk_attempt_id (attempt_id),
    KEY idx_transaction_id (transaction_id),
    KEY idx_original_transaction_id (original_transaction_id),
    KEY idx_uid (uid),
    KEY idx_app_account_token (app_account_token),
    KEY idx_client_status (client_status),
    KEY idx_server_status (server_status)
) COMMENT='iOS客户端购买尝试与同步状态表';
    """)

    add_heading(doc, "6.2 现有正式表扩字段 SQL", level=2)
    add_code_block(doc, """
ALTER TABLE t_user_subscription_contract
    ADD COLUMN app_account_token CHAR(36) DEFAULT NULL COMMENT '购买时对应的appAccountToken' AFTER product_id,
    ADD COLUMN anonymous_user_id VARCHAR(64) DEFAULT NULL COMMENT '匿名购买主体ID' AFTER app_account_token,
    ADD COLUMN latest_transaction_id VARCHAR(128) DEFAULT NULL COMMENT '最近一笔Apple transaction_id' AFTER contract_code,
    ADD COLUMN environment VARCHAR(16) DEFAULT NULL COMMENT 'Apple环境：Sandbox/Production' AFTER latest_transaction_id,
    ADD COLUMN last_notification_uuid VARCHAR(64) DEFAULT NULL COMMENT '最近一次驱动状态变化的Apple通知ID' AFTER environment,
    ADD COLUMN last_verify_time DATETIME DEFAULT NULL COMMENT '最近一次核验Apple成功时间' AFTER last_notification_uuid,
    ADD COLUMN last_error_msg VARCHAR(255) DEFAULT NULL COMMENT '最近一次异常原因' AFTER last_verify_time,
    ADD COLUMN recover_window_id BIGINT UNSIGNED DEFAULT NULL COMMENT '最近一次由哪次故障恢复任务补齐' AFTER last_error_msg;

ALTER TABLE t_user_subscription_order
    ADD COLUMN original_transaction_id VARCHAR(128) DEFAULT NULL COMMENT 'Apple自动续订链ID' AFTER contract_id,
    ADD COLUMN web_order_line_item_id VARCHAR(128) DEFAULT NULL COMMENT 'Apple续费行项目ID' AFTER original_transaction_id,
    ADD COLUMN app_account_token CHAR(36) DEFAULT NULL COMMENT '购买时对应的appAccountToken' AFTER web_order_line_item_id,
    ADD COLUMN anonymous_user_id VARCHAR(64) DEFAULT NULL COMMENT '匿名购买主体ID' AFTER app_account_token,
    ADD COLUMN source VARCHAR(16) DEFAULT NULL COMMENT '订单来源：client_report/notification/reconcile/manual' AFTER anonymous_user_id,
    ADD COLUMN verify_status VARCHAR(24) DEFAULT NULL COMMENT 'Apple校验状态：verified/unverified/retryable_failed' AFTER source,
    ADD COLUMN refund_reason VARCHAR(128) DEFAULT NULL COMMENT '退款原因说明' AFTER refund_time,
    ADD COLUMN recover_window_id BIGINT UNSIGNED DEFAULT NULL COMMENT '最近一次由哪次故障恢复任务补齐' AFTER refund_reason;

ALTER TABLE t_user_vip_info
    ADD COLUMN source VARCHAR(16) DEFAULT NULL COMMENT '会员来源：ios/wx/alipay/manual' AFTER status,
    ADD COLUMN source_contract_id VARCHAR(128) DEFAULT NULL COMMENT '当前生效权益对应的订阅链ID' AFTER source,
    ADD COLUMN sync_status VARCHAR(24) DEFAULT NULL COMMENT '会员同步状态：synced/pending/failed' AFTER source_contract_id,
    ADD COLUMN grace_status VARCHAR(24) DEFAULT NULL COMMENT '宽限与重试状态：none/grace/billing_retry' AFTER sync_status,
    ADD COLUMN last_verify_time DATETIME DEFAULT NULL COMMENT '最近一次核验会员态时间' AFTER grace_status,
    ADD COLUMN remark VARCHAR(255) DEFAULT NULL COMMENT '补偿或异常备注' AFTER last_verify_time,
    ADD COLUMN recover_window_id BIGINT UNSIGNED DEFAULT NULL COMMENT '最近一次由哪次故障恢复任务修正会员态' AFTER remark;
    """)

    add_heading(doc, "7. 接口设计", level=1)
    add_para(doc, "每个接口都分成：用途、为什么要有、请求字段、响应字段、服务端处理步骤。开发只要按这一节实现，就能串起完整链路。")

    interfaces = {
        "POST /users/apple/iap/identity/init": {
            "purpose": "注册前付费墙初始化匿名身份。",
            "why": "GuidanceProVC 在登录前就会购买，不能等登录后才有订阅身份。",
            "req": [
                {"name":"device_install_id","type":"string","required":"是","source":"iOS本地生成UUID","when":"请求时","comment":"设备安装实例ID，便于排查和幂等初始化。"},
                {"name":"app_version","type":"string","required":"是","source":"iOS Bundle","when":"请求时","comment":"App版本号，用于排查兼容性问题。"},
                {"name":"app_build","type":"string","required":"是","source":"iOS Bundle","when":"请求时","comment":"Build号。"},
            ],
            "resp": [
                {"name":"code","type":"int","required":"是","source":"服务端","when":"响应时","comment":"业务状态码，200表示成功。"},
                {"name":"message","type":"string","required":"是","source":"服务端","when":"响应时","comment":"提示信息。"},
                {"name":"anonymous_user_id","type":"string","required":"是","source":"服务端生成","when":"响应时","comment":"匿名购买主体ID。前端后续购买、恢复、注册绑定都要带它。"},
                {"name":"app_account_token","type":"string","required":"是","source":"服务端生成","when":"响应时","comment":"购买时传给Apple的appAccountToken。"},
                {"name":"expires_at","type":"string","required":"否","source":"服务端","when":"响应时","comment":"匿名身份有效期，可选。"},
                {"name":"server_time","type":"string","required":"是","source":"服务端","when":"响应时","comment":"服务端时间，用于前后端排障对齐。"},
            ],
            "steps": [
                "按 device_install_id 查询是否已有有效匿名身份。",
                "若已有且未绑定冲突，返回已有 anonymous_user_id + app_account_token。",
                "若没有，则新建 t_ios_identity。",
                "返回匿名身份给客户端。",
            ],
        },
        "POST /users/apple/iap/purchase/report": {
            "purpose": "客户端 verified transaction 后立即上报交易。",
            "why": "不要再只依赖 users/apple/iap/order/query 去做所有事情；这里是正式的交易上报入口。",
            "req": [
                {"name":"attempt_id","type":"string","required":"是","source":"前端UUID","when":"请求时","comment":"一次购买尝试ID。"},
                {"name":"uid","type":"string","required":"否","source":"当前登录用户","when":"请求时","comment":"已登录购买时传；注册前购买可为空。"},
                {"name":"anonymous_user_id","type":"string","required":"否","source":"identity/init","when":"请求时","comment":"注册前购买场景必须传。"},
                {"name":"app_account_token","type":"string","required":"否","source":"identity/init","when":"请求时","comment":"购买时传给Apple的身份token。"},
                {"name":"transaction_id","type":"string","required":"是","source":"Apple verified transaction","when":"请求时","comment":"Apple单笔交易ID。"},
                {"name":"original_transaction_id","type":"string","required":"是","source":"Apple verified transaction","when":"请求时","comment":"Apple自动续订链ID。"},
                {"name":"product_id","type":"string","required":"是","source":"Apple verified transaction","when":"请求时","comment":"商品ID。"},
                {"name":"purchase_date_ms","type":"long","required":"是","source":"Apple verified transaction","when":"请求时","comment":"购买时间，毫秒时间戳。"},
                {"name":"signed_transaction_info","type":"string","required":"否","source":"StoreKit 2 verified transaction","when":"请求时","comment":"原始 signedTransactionInfo JWS。"},
                {"name":"signed_renewal_info","type":"string","required":"否","source":"StoreKit 2 entitlement/renewal info","when":"请求时","comment":"原始 signedRenewalInfo JWS。"},
                {"name":"environment","type":"string","required":"是","source":"Apple / 客户端","when":"请求时","comment":"Sandbox 或 Production。"},
                {"name":"device_install_id","type":"string","required":"否","source":"iOS本地","when":"请求时","comment":"设备安装实例ID。"},
                {"name":"client_status","type":"string","required":"是","source":"前端","when":"请求时","comment":"客户端状态，如 verified_local、pending_sync。"},
                {"name":"app_version","type":"string","required":"是","source":"iOS Bundle","when":"请求时","comment":"App版本。"},
                {"name":"app_build","type":"string","required":"是","source":"iOS Bundle","when":"请求时","comment":"Build号。"},
            ],
            "resp": [
                {"name":"code","type":"int","required":"是","source":"服务端","when":"响应时","comment":"业务状态码。"},
                {"name":"message","type":"string","required":"是","source":"服务端","when":"响应时","comment":"提示文案。"},
                {"name":"accept_status","type":"string","required":"是","source":"服务端","when":"响应时","comment":"verified / pending / retryable_failed / rejected。"},
                {"name":"membership_status","type":"string","required":"是","source":"服务端","when":"响应时","comment":"当前会员状态：active/expired/pending/grace/billing_retry等。"},
                {"name":"vip_type","type":"int","required":"否","source":"服务端","when":"响应时","comment":"0非会员，1月，2年，3终身。"},
                {"name":"expire_time","type":"string","required":"否","source":"服务端","when":"响应时","comment":"当前会员到期时间。"},
                {"name":"need_bind_login","type":"boolean","required":"是","source":"服务端","when":"响应时","comment":"若为 true，表示交易已记账但尚未归属到真实 uid。"},
                {"name":"next_retry_after_sec","type":"int","required":"否","source":"服务端","when":"响应时","comment":"【本次新增-重启容灾】若后台刚恢复或下游依赖尚未稳定，可返回建议客户端下次重试秒数，避免重启后重试风暴。"},
                {"name":"server_trace_id","type":"string","required":"是","source":"服务端","when":"响应时","comment":"服务端日志追踪ID。"},
            ],
            "steps": [
                "先按 attempt_id upsert t_ios_purchase_attempt。",
                "按 transaction_id 查 ledger，若已存在则走幂等返回。",
                "验证 JWS；必要时调用 Get Transaction Info / Get All Subscription Statuses。",
                "写 t_ios_transaction_ledger。",
                "若 uid 已知，则同步更新正式三表。",
                "若 uid 未知，则只挂匿名身份，返回 need_bind_login=true。",
                "【本次修改-重启容灾】若交易已成功落 ledger 但正式三表因重启或数据库抖动未补齐，则 formal_status 标记为 reconcile_pending，等待恢复任务自动补齐。",
            ],
        },
        "POST /users/apple/iap/order/query": {
            "purpose": "兼容现有客户端的交易查单与补绑接口。",
            "why": "当前 iOS 代码已经大量使用该接口，短期内不宜直接废弃。",
            "req": [
                {"name":"transactionId","type":"string","required":"是","source":"现有客户端","when":"请求时","comment":"Apple单笔交易ID。当前 iOS 已在用。"},
                {"name":"bizType","type":"string","required":"是","source":"现有客户端","when":"请求时","comment":"业务场景：1 pending bind，2 AI guidance，3 standard。"},
                {"name":"uid","type":"string","required":"否","source":"当前登录用户","when":"建议新增并传入","comment":"若当前已登录，便于服务端直接做匿名绑定。"},
                {"name":"anonymous_user_id","type":"string","required":"否","source":"identity/init","when":"建议新增并传入","comment":"注册前购买绑定辅助字段。"},
                {"name":"app_account_token","type":"string","required":"否","source":"identity/init","when":"建议新增并传入","comment":"辅助查 ledger 和 identity。"},
            ],
            "resp": [
                {"name":"code","type":"int","required":"是","source":"服务端","when":"响应时","comment":"200 表示查单或补绑成功。"},
                {"name":"message","type":"string","required":"是","source":"服务端","when":"响应时","comment":"响应说明。"},
                {"name":"membership_status","type":"string","required":"否","source":"服务端","when":"响应时","comment":"当前会员态。"},
                {"name":"vip_type","type":"int","required":"否","source":"服务端","when":"响应时","comment":"当前会员类型。"},
                {"name":"expire_time","type":"string","required":"否","source":"服务端","when":"响应时","comment":"到期时间。"},
                {"name":"bound_uid","type":"string","required":"否","source":"服务端","when":"响应时","comment":"交易最终归属到的 uid。"},
                {"name":"server_trace_id","type":"string","required":"是","source":"服务端","when":"响应时","comment":"日志追踪ID。"},
            ],
            "steps": [
                "按 transactionId 查 ledger。",
                "若 ledger 不存在，可尝试从 Apple 再查一次；仍不存在则返回未检测到交易。",
                "若 ledger 存在且当前请求带 uid，且该交易仍未正式归属，则执行绑定和正式表补齐。",
                "若正式表已齐全，则直接幂等返回当前会员状态。",
                "【本次修改-重启容灾】若发现 ledger 已存在、正式表缺失，说明可能是后台重启导致上次响应中断，本次请求必须直接走补齐，而不是返回未找到。",
            ],
        },
        "POST /users/apple/iap/bind-anonymous": {
            "purpose": "注册或登录成功后，把匿名购买绑定到真实账号。",
            "why": "Elavatine 存在注册前购买，这是匿名身份转正式账号的关键接口。",
            "req": [
                {"name":"uid","type":"string","required":"是","source":"登录用户","when":"请求时","comment":"真实业务用户 ID。"},
                {"name":"anonymous_user_id","type":"string","required":"是","source":"identity/init","when":"请求时","comment":"匿名主体。"},
                {"name":"app_account_token","type":"string","required":"是","source":"identity/init","when":"请求时","comment":"购买身份 token。"},
                {"name":"device_install_id","type":"string","required":"否","source":"iOS本地","when":"请求时","comment":"设备安装实例ID。"},
            ],
            "resp": [
                {"name":"code","type":"int","required":"是","source":"服务端","when":"响应时","comment":"业务状态码。"},
                {"name":"message","type":"string","required":"是","source":"服务端","when":"响应时","comment":"说明文案。"},
                {"name":"bind_status","type":"string","required":"是","source":"服务端","when":"响应时","comment":"success / conflict / no_purchase / pending。"},
                {"name":"membership_status","type":"string","required":"否","source":"服务端","when":"响应时","comment":"绑定后的会员状态。"},
                {"name":"vip_type","type":"int","required":"否","source":"服务端","when":"响应时","comment":"绑定后的会员类型。"},
                {"name":"expire_time","type":"string","required":"否","source":"服务端","when":"响应时","comment":"绑定后的到期时间。"},
                {"name":"recovered_from_outage","type":"boolean","required":"否","source":"服务端","when":"响应时","comment":"【本次新增-重启容灾】若本次绑定顺带补齐了重启窗口内缺失的正式表，则返回 true，便于前端与客服埋点。"},
            ],
            "steps": [
                "查 t_ios_identity，确认 anonymous_user_id + app_account_token 存在。",
                "查该匿名身份下是否有有效 ledger。",
                "若无有效交易，返回 no_purchase。",
                "若该匿名身份已绑定其他 uid，返回 conflict，并记录告警。",
                "若可以绑定，则更新 t_ios_identity.uid + bind_status。",
                "回放该身份下未正式落账的交易，补齐 contract/order/vip_info。",
            ],
        },
        "POST /users/apple/iap/restore": {
            "purpose": "恢复购买后，把客户端发现的有效 entitlement 重放给后台。",
            "why": "恢复购买是 iOS 标准补单入口，不能只靠用户重新购买。",
            "req": [
                {"name":"uid","type":"string","required":"是","source":"登录用户","when":"请求时","comment":"当前用户ID。"},
                {"name":"app_account_token","type":"string","required":"否","source":"identity/init 或用户身份","when":"请求时","comment":"当前账号对应 token。"},
                {"name":"transactions","type":"array","required":"是","source":"Transaction.currentEntitlements","when":"请求时","comment":"客户端扫描到的有效交易数组。"},
                {"name":"transactions[].transaction_id","type":"string","required":"是","source":"Apple entitlement","when":"请求时","comment":"单笔交易ID。"},
                {"name":"transactions[].original_transaction_id","type":"string","required":"是","source":"Apple entitlement","when":"请求时","comment":"订阅链ID。"},
                {"name":"transactions[].product_id","type":"string","required":"是","source":"Apple entitlement","when":"请求时","comment":"商品ID。"},
                {"name":"transactions[].signed_transaction_info","type":"string","required":"否","source":"StoreKit 2","when":"请求时","comment":"原始 signedTransactionInfo。"},
                {"name":"transactions[].signed_renewal_info","type":"string","required":"否","source":"StoreKit 2","when":"请求时","comment":"原始 signedRenewalInfo。"},
                {"name":"transactions[].environment","type":"string","required":"是","source":"Apple","when":"请求时","comment":"Sandbox/Production。"},
            ],
            "resp": [
                {"name":"code","type":"int","required":"是","source":"服务端","when":"响应时","comment":"业务状态码。"},
                {"name":"message","type":"string","required":"是","source":"服务端","when":"响应时","comment":"响应说明。"},
                {"name":"restored_count","type":"int","required":"是","source":"服务端","when":"响应时","comment":"成功恢复并补齐到后台的交易数量。"},
                {"name":"membership_status","type":"string","required":"是","source":"服务端","when":"响应时","comment":"恢复后的会员状态。"},
                {"name":"vip_type","type":"int","required":"否","source":"服务端","when":"响应时","comment":"恢复后的会员类型。"},
                {"name":"expire_time","type":"string","required":"否","source":"服务端","when":"响应时","comment":"恢复后的到期时间。"},
            ],
            "steps": [
                "逐笔遍历 transactions。",
                "按 transaction_id 幂等写入或补全 ledger。",
                "若未归属当前 uid，则按 original_transaction_id 和 app_account_token 尝试绑定。",
                "补齐正式三表并返回当前会员态。",
            ],
        },
        "POST /users/apple/iap/notifications": {
            "purpose": "接收 Apple App Store Server Notifications V2。",
            "why": "续费、退款、撤销、宽限期等事件不能只靠客户端知道。",
            "req": [
                {"name":"signedPayload","type":"string","required":"是","source":"Apple","when":"回调请求","comment":"Apple 原始通知 signedPayload JWS。"},
            ],
            "resp": [
                {"name":"http_status","type":"int","required":"是","source":"服务端","when":"HTTP响应","comment":"成功接收并入队后返回 200-206，避免 Apple 重试。"},
                {"name":"body","type":"string","required":"否","source":"服务端","when":"HTTP响应","comment":"通常可为空或返回简单 OK。"},
            ],
            "steps": [
                "先把 signedPayload 原样写入 t_ios_notification_inbox，状态为 pending。",
                "校验 signedPayload JWS。",
                "将通知投递到异步处理器。",
                "【本次修改-重启容灾】只有在 inbox insert 已经成功提交事务后，才允许返回 HTTP 200；如果服务在落库前重启或数据库失败，必须返回非 2xx，让 Apple 后续重试。",
                "成功入队即返回 HTTP 200。",
                "异步处理器中再更新 ledger 和正式三表。",
            ],
        },
        "GET /users/membership/vip_info": {
            "purpose": "前端唯一正式会员状态源。",
            "why": "大量现有页面依赖该接口，不建议替换，只建议增强字段。",
            "req": [
                {"name":"Authorization / token","type":"header","required":"是","source":"当前登录用户","when":"请求时","comment":"用于确定 uid。"},
            ],
            "resp": [
                {"name":"uid","type":"string","required":"是","source":"t_user_vip_info","when":"响应时","comment":"用户ID。"},
                {"name":"vip_type","type":"int","required":"是","source":"t_user_vip_info","when":"响应时","comment":"0非会员，1月，2年，3终身。"},
                {"name":"is_lifetime","type":"int","required":"是","source":"t_user_vip_info","when":"响应时","comment":"是否终身会员。"},
                {"name":"start_time","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"当前周期开始时间。"},
                {"name":"expire_time","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"当前周期结束时间。"},
                {"name":"status","type":"int","required":"是","source":"t_user_vip_info","when":"响应时","comment":"1有效，2过期，3冻结。"},
                {"name":"source","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"当前会员来源。"},
                {"name":"source_contract_id","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"当前生效权益对应订阅链。"},
                {"name":"sync_status","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"synced/pending/failed，前端可用来显示“订单同步中”。"},
                {"name":"grace_status","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"none/grace/billing_retry。"},
                {"name":"last_verify_time","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"最近一次会员态核验时间。"},
                {"name":"remark","type":"string","required":"否","source":"t_user_vip_info","when":"响应时","comment":"备注。"},
            ],
            "steps": [
                "按 uid 读取 t_user_vip_info。",
                "若发现 sync_status=pending，也可额外查 contract/order 返回更精确文案。",
                "该接口只返回正式业务状态，不返回原始 ledger。"
            ],
        },
    }

    for name, meta in interfaces.items():
        add_heading(doc, name, level=2)
        add_para(doc, f"用途：{meta['purpose']}")
        add_para(doc, f"为什么要有：{meta['why']}")
        add_field_table(doc, f"{name} 请求字段", meta["req"])
        add_field_table(doc, f"{name} 响应字段", meta["resp"])
        add_para(doc, "服务端处理步骤：")
        add_numbered(doc, meta["steps"])

    add_heading(doc, "8. Apple 回调处理规则", level=1)
    add_para(doc, "本节解决“Apple 通知来了以后后台到底怎么处理”。研发只要按本节的规则写异步消费器，就不会把通知逻辑写散。")
    add_numbered(doc, [
        "收到 Apple 回调后，先把 signedPayload 原样写入 t_ios_notification_inbox，process_status = pending。",
        "读取 notificationUUID，若 inbox 中已存在同 UUID，则直接返回 HTTP 200，避免重复处理。",
        "校验 signedPayload JWS，解析出 notificationType、subtype、signedTransactionInfo、signedRenewalInfo。",
        "若解析失败，更新 inbox.process_status = failed，并记录 process_error；但仍然保留原始 payload。",
        "若解析成功，则写入或补全 t_ios_transaction_ledger。",
        "再根据交易链和状态去更新正式业务三表。",
    ])
    add_marker_note(doc, MARKER_MOD, "这里最关键的容灾要求是：通知必须先入 inbox 再返回 200。否则服务刚好在返回前重启，就会出现 Apple 以为你收到了、但后台实际上没落库的最坏情况。")

    add_heading(doc, "8.1 各种 Apple 通知场景如何处理", level=2)
    add_table(doc,
              ["场景", "如何识别", "ledger 怎么写", "正式三表怎么写", "为什么这么做"],
              [
                  ["首购成功", "SUBSCRIBED 或首个 verified 交易", "新增 ledger，status=verified", "若已绑定 uid，则建 contract/order/vip_info；否则仅挂匿名身份", "注册前购买不能直接写正式 uid 表。"],
                  ["续费成功", "DID_RENEW", "新增一条 ledger", "新增 order，更新 contract.latest_transaction_id/next_deduct_time，更新 vip_info.expire_time", "续费是新交易，不是覆盖旧交易。"],
                  ["用户关闭自动续费", "DID_CHANGE_RENEWAL_STATUS 且 auto renew off", "可补写 renewal 状态", "contract.status=2，写 cancel_time，vip_info 仍保持有效直到实际过期", "关闭续费不等于立刻失效。"],
                  ["账单重试", "DID_FAIL_TO_RENEW", "ledger_status=billing_retry", "contract.status=3，vip_info.grace_status=billing_retry", "保留恢复空间，不直接判失败。"],
                  ["宽限期", "GRACE_PERIOD_EXPIRED 前的状态或 renewal info", "ledger_status=grace", "vip_info.grace_status=grace", "前端和产品可决定是否继续给权益。建议继续给。"],
                  ["过期", "EXPIRED", "ledger_status=expired", "contract.status=4 或过期态，vip_info.status=2", "正式降权。"],
                  ["退款", "REFUND", "ledger_status=refunded", "order.status=3，refund_time 写入，vip_info 降权", "退款必须撤销权益。"],
                  ["撤销", "REVOKE", "ledger_status=revoked", "vip_info 立即降权", "撤销通常比退款更强。"],
                  ["升级/降级", "DID_CHANGE_RENEWAL_PREF / 新 productId", "写新 ledger", "contract.product_id 更新，vip_info.vip_type/expire_time 更新", "订阅链不变，但产品可能变。"],
              ])

    add_heading(doc, "9. 幂等与防错规则", level=1)
    add_bullets(doc, [
        "同一 transaction_id 重复上报，不能重复开订单，也不能报错导致客户端误判失败。",
        "同一 notification_uuid 重复回调，只能处理一次。",
        "正式合同表以 (contract_id, platform) 唯一约束防止同一订阅链多开。",
        "正式订单表以 third_party_trade_no 唯一约束防止同一 transaction_id 多开。",
        "匿名身份首次绑定到 uid 后，默认冻结，后续不同 uid 冲突时返回 conflict 并告警，不自动改绑。",
    ])

    add_heading(doc, "10. 服务器重启/不可用场景处理", level=1)
    add_marker_note(doc, MARKER_NEW, "本节是本次重点补充。目标是覆盖：首购时后台重启、登录后绑定时后台重启、Apple 续费/退款通知投递时后台重启。")
    add_heading(doc, "10.1 必须防住的四类故障窗口", level=2)
    add_table(doc,
              ["故障窗口", "可能丢什么", "如果不处理会怎样", "完整版方案怎么兜"],
              [
                  ["首购时后台重启", "purchase/report 请求未落账", "Apple 已扣费，但后台看不到交易", "客户端 pending queue 重试 + ledger 幂等入账 + restore + 对账任务"],
                  ["登录后绑定时后台重启", "匿名购买未补到正式三表", "用户本地可用，后台正式会员态缺失", "bind-anonymous / order-query 幂等补齐 + formal_status=reconcile_pending + 恢复任务补齐"],
                  ["Apple 续费通知时后台重启", "DID_RENEW / DID_CHANGE_RENEWAL_STATUS 等通知", "续费成功但会员到期时间没续上，或取消状态不同步", "ASSN V2 重试 + Notification History 回补 + 订阅状态对账"],
                  ["Apple 退款/撤销通知时后台重启", "REFUND / REVOKE 等通知", "该降权却没降，账实不一致", "ASSN V2 重试 + Notification History 回补 + 主动状态校验"],
              ])

    add_heading(doc, "10.2 后台接口在重启窗口内的行为要求", level=2)
    add_bullets(doc, [
        "purchase/report：若请求在服务重启前已写入 ledger、但响应丢失，客户端重试时必须按 transaction_id 返回幂等成功，而不是重新开单或报错。",
        "order/query：若发现 ledger 已存在、正式三表缺失，必须直接补齐正式表，不能简单返回“未查到订单”。",
        "bind-anonymous：若上次绑定做到一半就重启，本次重试要以 anonymous_user_id/app_account_token 为主键继续补齐，而不是重新创建一套记录。",
        "notifications：只有在 inbox 落库事务提交成功后，才允许返回 200。",
    ])

    add_heading(doc, "10.3 自动恢复链路", level=2)
    add_numbered(doc, [
        "客户端恢复链：客户端 verified 后若上报超时，本地 pending queue 保留交易，等待下次启动、登录、前台激活、恢复购买时继续重试。",
        "Apple 恢复链：ASSN V2 在生产环境会重试；后台重启恢复后，仍要继续接收后续重试通知。",
        "服务端恢复链：后台重启恢复后，根据 t_ios_outage_window 跑恢复任务，补拉 Notification History，补扫 ledger 与正式三表不一致的数据。",
        "人工恢复链：若系统自动恢复后仍有 conflict/no_purchase 等特殊状态，再交给客服与技术人工排查。",
    ])

    add_heading(doc, "10.4 故障恢复任务的执行顺序", level=2)
    add_numbered(doc, [
        "先根据部署/健康检查/实例重启日志创建或关闭 t_ios_outage_window。",
        "恢复后先扫该窗口内的 t_ios_notification_inbox 缺口，调用 Get Notification History 回补通知。",
        "再扫该窗口内 ledger.formal_status in (reconcile_pending,reconcile_failed) 的交易，补齐正式三表。",
        "再扫 purchase_attempt.server_ack_status in (none,response_lost) 且 transaction_id 不为空的购买尝试，按 transaction_id 查 ledger 并修正 server_status。",
        "最后按 original_transaction_id 对最近活跃订阅链抽样或全量调用 Get All Subscription Statuses，确保续费、宽限期、退款状态和正式表一致。",
    ])

    add_heading(doc, "10.5 Notification History 回补规则", level=2)
    add_bullets(doc, [
        "只要判定 notifications 入口在某次故障窗口内不可用，就必须跑 Get Notification History，而不是只等 Apple 自己重试。",
        "回补到的通知写入 t_ios_notification_inbox 时，is_replayed=1。",
        "如果 Notification History 返回 firstSendAttemptResult=NO_RESPONSE/TIMED_OUT，可直接证明这条通知是在服务不可用期间漏掉的。",
        "回补通知后仍按正常异步消费器处理，不允许写旁路逻辑直接跳过 inbox。",
    ])

    add_heading(doc, "11. 对账任务设计", level=1)
    add_para(doc, "对账任务不是锦上添花，而是成熟订阅系统的必要组成。因为客户端链路、通知链路、后台链路都可能临时失败。")
    add_bullets(doc, [
        "任务A：扫 t_ios_purchase_attempt 中 server_status=retryable_failed 或 client_status=pending_sync 的记录。",
        "任务B：扫 t_ios_notification_inbox 中 process_status=failed 的通知，重放处理。",
        "任务C：扫 t_ios_transaction_ledger 中 ledger 已 verified 但正式三表未落齐的记录，自动补齐。",
        "任务D：扫 contract.status=3 或 vip_info.grace_status in (grace,billing_retry) 的订阅，主动查 Apple 最新状态。",
        "任务E：扫注册前购买后长时间未绑定的匿名身份，输出给客服或风控排查。",
    ])
    add_para(doc, "服务端主动查 Apple 时，优先使用：Get Transaction Info、Get Transaction History、Get All Subscription Statuses；怀疑漏通知时再用 Get Notification History。")
    add_marker_note(doc, MARKER_MOD, "对账任务现在不只用于日常查漏，还要显式支持“按故障窗口回放”的模式。也就是先定窗口，再定链路，再回补。")

    add_heading(doc, "12. 与当前 iOS 代码的对齐说明", level=1)
    add_bullets(doc, [
        "当前 iOS 端已有 users/apple/iap/order/query，可先保留作为兼容补绑入口，但建议新增 purchase/report 作为正式交易上报入口。",
        "当前 GuidanceProVC 是注册前购买，这就是为什么 identity/init 与 bind-anonymous 是必需接口。",
        "当前 vip_info 仍是前端正式会员真相，因此后台改造后仍要继续维护 t_user_vip_info。",
        "当前客户端已有 pending transaction 雏形，但字段不足；后台的 t_ios_purchase_attempt 正好承接前端未来的 pending queue。"
    ])

    add_heading(doc, "13. 上线顺序建议", level=1)
    add_numbered(doc, [
        "先建表：t_ios_identity、t_ios_transaction_ledger、t_ios_notification_inbox、t_ios_purchase_attempt。",
        "【本次新增-重启容灾】同步建 t_ios_outage_window。",
        "再扩正式三表。",
        "先接 App Store Server Notifications V2 收件与入库，不急着一次性做完所有业务处理。",
        "实现 identity/init 与 purchase/report。",
        "实现 bind-anonymous。",
        "增强 users/apple/iap/order/query 为兼容补单入口。",
        "实现 restore。",
        "【本次新增-重启容灾】最后补 Notification History 回补与按故障窗口执行的恢复任务。",
    ])

    add_heading(doc, "14. 一句话给研发的原则", level=1)
    add_para(doc, "不要把订阅当成一次同步支付请求，而要把它当成一套 Apple 交易驱动、客户端可恢复、后台可幂等、正式业务表可自动补齐的账本系统。这样才有可能真正避免“用户已付款但会员未生效”的事故。")

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
